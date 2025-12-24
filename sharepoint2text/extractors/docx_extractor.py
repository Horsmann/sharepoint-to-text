"""
DOCX content extractor using python-docx library.
"""

import datetime
import io
import logging

from docx import Document
from docx.oxml.ns import qn

logging.basicConfig(level=logging.DEBUG)

logger = logging.getLogger(__name__)


def read_docx(file_like: io.BytesIO) -> dict:
    """
    Extract all relevant content from a DOCX file.

    Args:
        file_like: A BytesIO object containing the DOCX file data.

    Returns:
        Dictionary with all extracted content including images as BytesIO objects.
    """
    doc = Document(file_like)
    result = {}

    # === Core Properties (Metadata) ===
    props = doc.core_properties
    result["metadata"] = {
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "category": props.category or "",
        "comments": props.comments or "",
        "created": props.created.isoformat()
        if isinstance(props.created, datetime.datetime)
        else "",
        "modified": props.modified.isoformat()
        if isinstance(props.modified, datetime.datetime)
        else "",
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision,
    }

    result["paragraphs"] = []
    for para in doc.paragraphs:
        para_data = {
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": [],
        }
        for run in para.runs:
            para_data["runs"].append(
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_color": str(run.font.color.rgb)
                    if run.font.color and run.font.color.rgb
                    else None,
                }
            )
        result["paragraphs"].append(para_data)

    result["tables"] = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                row_data.append(cell_text)
            table_data.append(row_data)
        result["tables"].append(table_data)

    result["headers"] = []
    result["footers"] = []
    for section in doc.sections:
        # Default
        if section.header and section.header.paragraphs:
            text = "\n".join(p.text for p in section.header.paragraphs)
            if text.strip():
                result["headers"].append({"type": "default", "text": text})
        if section.footer and section.footer.paragraphs:
            text = "\n".join(p.text for p in section.footer.paragraphs)
            if text.strip():
                result["footers"].append({"type": "default", "text": text})

        # First page
        if section.first_page_header and section.first_page_header.paragraphs:
            text = "\n".join(p.text for p in section.first_page_header.paragraphs)
            if text.strip():
                result["headers"].append({"type": "first_page", "text": text})
        if section.first_page_footer and section.first_page_footer.paragraphs:
            text = "\n".join(p.text for p in section.first_page_footer.paragraphs)
            if text.strip():
                result["footers"].append({"type": "first_page", "text": text})

        # Even page
        if section.even_page_header and section.even_page_header.paragraphs:
            text = "\n".join(p.text for p in section.even_page_header.paragraphs)
            if text.strip():
                result["headers"].append({"type": "even_page", "text": text})
        if section.even_page_footer and section.even_page_footer.paragraphs:
            text = "\n".join(p.text for p in section.even_page_footer.paragraphs)
            if text.strip():
                result["footers"].append({"type": "even_page", "text": text})

    result["images"] = []
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                result["images"].append(
                    {
                        "rel_id": rel_id,
                        "filename": image_part.partname.split("/")[-1],
                        "content_type": image_part.content_type,
                        "data": io.BytesIO(image_part.blob),
                        "size_bytes": len(image_part.blob),
                    }
                )
            except Exception as e:
                logger.debug(f"Image extraction failed for rel_id {rel_id} - {e}")
                result["images"].append({"rel_id": rel_id, "error": str(e)})

    result["hyperlinks"] = []
    rels = doc.part.rels
    for para in doc.paragraphs:
        for hyperlink in para._element.findall(
            ".//w:hyperlink",
            {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        ):
            r_id = hyperlink.get(qn("r:id"))
            if r_id and r_id in rels and "hyperlink" in rels[r_id].reltype:
                text = "".join(
                    t.text or ""
                    for t in hyperlink.findall(
                        ".//w:t",
                        {
                            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        },
                    )
                )
                result["hyperlinks"].append(
                    {"text": text, "url": rels[r_id].target_ref}
                )

    result["footnotes"] = []
    try:
        if doc.part.footnotes_part:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for fn in doc.part.footnotes_part.element.findall(".//w:footnote", ns):
                fn_id = fn.get(qn("w:id"))
                if fn_id not in ["-1", "0"]:
                    text = "".join(t.text or "" for t in fn.findall(".//w:t", ns))
                    result["footnotes"].append({"id": fn_id, "text": text})
    except AttributeError as e:
        logger.debug(f"Silently ignoring footnote extraction error {e}")
        pass

    result["endnotes"] = []
    try:
        if doc.part.endnotes_part:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for en in doc.part.endnotes_part.element.findall(".//w:endnote", ns):
                en_id = en.get(qn("w:id"))
                if en_id not in ["-1", "0"]:
                    text = "".join(t.text or "" for t in en.findall(".//w:t", ns))
                    result["endnotes"].append({"id": en_id, "text": text})
    except AttributeError as e:
        logger.debug(f"Silently ignoring endnote extraction error {e}")
        pass

    # === Comments ===
    result["comments"] = []
    try:
        if doc.part.comments_part:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for comment in doc.part.comments_part.element.findall(".//w:comment", ns):
                result["comments"].append(
                    {
                        "id": comment.get(qn("w:id")),
                        "author": comment.get(qn("w:author")),
                        "date": comment.get(qn("w:date")),
                        "text": "".join(
                            t.text or "" for t in comment.findall(".//w:t", ns)
                        ),
                    }
                )
    except AttributeError as e:
        logger.debug(f"Silently ignoring comments extraction error {e}")
        pass

    # === Sections (page layout) ===
    result["sections"] = []
    for section in doc.sections:
        result["sections"].append(
            {
                "page_width_inches": section.page_width.inches
                if section.page_width
                else None,
                "page_height_inches": section.page_height.inches
                if section.page_height
                else None,
                "left_margin_inches": section.left_margin.inches
                if section.left_margin
                else None,
                "right_margin_inches": section.right_margin.inches
                if section.right_margin
                else None,
                "top_margin_inches": section.top_margin.inches
                if section.top_margin
                else None,
                "bottom_margin_inches": section.bottom_margin.inches
                if section.bottom_margin
                else None,
                "orientation": str(section.orientation)
                if section.orientation
                else None,
            }
        )

    # === Styles used ===
    styles = set()
    for para in doc.paragraphs:
        if para.style:
            styles.add(para.style.name)
    result["styles"] = list(styles)

    # === Full text (convenience) ===
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                if text:
                    all_text.append(text)
    result["full_text"] = "\n".join(all_text)

    return result
